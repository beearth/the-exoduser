# -*- coding: utf-8 -*-
"""HELL: EXODUSER 최종기획서 생성 스크립트 — 기존 에셋 활용 버전"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

IMG_DIR = r'G:\hell\docs\최종기획서\images'
GAME_IMG = r'G:\hell\img'
SS_DIR = r'c:\Users\심도진\Pictures\Screenshots'
OUT_PATH = r'C:\Users\심도진\Desktop\HELL_EXODUSER_최종기획서_v4.1b.docx'

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00) if level == 1 else RGBColor(0x4A, 0x00, 0x00)
        run.font.name = 'Malgun Gothic'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return h

def add_para(text, bold=False, size=None, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if align: p.alignment = align
    return p

def add_img(path, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def img_game(name): return os.path.join(GAME_IMG, name)
def img_new(name): return os.path.join(IMG_DIR, name)
def img_ss(name): return os.path.join(SS_DIR, name)

def add_table_simple(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    run.font.size = Pt(8)

# ════════════════════════════════════════
# 표지 — FDG 로고 + EXODUSER 포스터
# ════════════════════════════════════════
doc.add_paragraph()
add_img(img_ss('스크린샷 2026-05-10 163644.png'), Inches(5))
doc.add_paragraph()
add_img(img_game('etype84_lightning.png'), Inches(4.5))
doc.add_paragraph()
add_para('GAME DESIGN DOCUMENT v4.1', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x66, 0x66, 0x66))
add_para('Dark Fantasy Hack & Slash ARPG  |  Solo Indie + AI-First  |  PC / Steam', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x66, 0x66, 0x66))
doc.add_paragraph()
add_para('FDG (ForDearGamers)', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
add_para('2026', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
doc.add_page_break()

# ════════════════════════════════════════
# 목차
# ════════════════════════════════════════
add_heading_styled('목차', 1)
toc_items = [
    '01. 한 페이지 개요 (Executive Summary)',
    '02. 세계관 & 스토리',
    '03. 아트 디렉션',
    '04. 핵심 캐릭터',
    '05. 월드 디자인 (7 지옥)',
    '06. 액션 정체성: Bullet Language',
    '07. 액션 정체성: Chain Drive',
    '08. 스킬 & 합체 시스템',
    '09. 아이템 & 어픽스',
    '10. 몬스터 디자인',
    '11. 보스 시스템',
    '12. 기술 아키텍처',
    '13. AI-First 파이프라인',
    '14. 개발 현황',
    '15. 출시 로드맵',
    '16. DLC & 포스트런칭',
]
for item in toc_items:
    add_para(item, size=11)
doc.add_page_break()

# ════════════════════════════════════════
# 01. EXECUTIVE SUMMARY
# ════════════════════════════════════════
add_heading_styled('01. 한 페이지 개요', 1)

add_para('ELEVATOR PITCH', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_para('복수를 끝낸 남자가 지옥에 떨어진다. 색으로 구분되는 탄막을 정확한 키로 받아치고, 사슬을 던져 벽에 박고 무적으로 끌려가며, 왕창 달려드는 괴물을 쓸어죽이는 2D 탑다운 ARPG.', size=11)
doc.add_paragraph()

# 실제 게임플레이 스크린샷
add_img(img_ss('스크린샷 2026-05-10 162633.png'), Inches(5.5))
doc.add_paragraph()

add_table_simple(
    ['지표', '수치'],
    [
        ['코드 규모', '30,000+ 줄 구현'],
        ['프레임레이트', '240 FPS 안정'],
        ['몬스터 종류', '100종'],
        ['보스 패턴', '48 패턴 / 58 콤보'],
        ['스킬', '22 액티브 + 11 패시브 + 21 합체'],
        ['어픽스', '55+ 종'],
        ['챕터', '7 챕터 / 35 스테이지 / 19 보스'],
    ]
)
doc.add_paragraph()
add_para('3대 핵심 필러 (Three Pillars)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_table_simple(
    ['필러', '설명'],
    [
        ['BULLET LANGUAGE', '색 x 키 매칭 패링. 잘못 치면 눈앞 폭발.'],
        ['CHAIN DRIVE', '홀딩 차지 사슬. 무적 돌진 + 콤보 캔슬.'],
        ['HELLSWARM', '어택 티켓 없음. 40+ 몹 동시 타격 240fps.'],
    ]
)
doc.add_paragraph()
add_para('사업 로드맵', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_para('BIC 2026 제출 -> Neowiz 퍼블리싱 피치 -> Steam Early Access $18-20 -> 풀릴리즈 $35-40 -> DLC $10-12', size=10)
doc.add_page_break()

# ════════════════════════════════════════
# 02. 세계관 & 스토리
# ════════════════════════════════════════
add_heading_styled('02. 세계관 & 스토리', 1)
add_para('"살아남은 자에게 주어진 것은, 오직 복수뿐."', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()

add_heading_styled('세계관 설정', 2)
add_para('지옥이란 무엇인가', bold=True, size=11)
add_para('지옥은 단순한 형벌의 공간이 아니다. 인간 세상에서 흘러내려온 악의(惡意)와 원한이 지하에 축적되어 태양처럼 세계를 지탱하는 에너지원이다. 지옥왕은 이 막대한 에너지를 지배하는 자이며, 그 권력은 신들의 세계에서도 최상위에 해당한다.')
add_para('신들의 세계란 무엇인가', bold=True, size=11)
add_para('신들의 세계는 성스럽고 이상적인 곳이 아니다. 인간 세상과 마찬가지로 정치, 음모, 권력 다툼이 존재한다. 지옥왕의 자리는 신들의 세계에서 가장 탐나는 권좌 중 하나다.')

add_separator()
add_heading_styled('스토리 구조', 2)

# ACT 1 — 전쟁 + 폐허 + 복수
add_para('ACT 1 — 지상: 복수 (Chapter 1)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_img(img_game('cin_war.png'), Inches(5))
add_para('오랜 전쟁에서 살아 돌아온 주인공. 그러나 귀환한 고향에는 폐허만이 남아있다.')
add_img(img_game('cin_ruins.png'), Inches(3.5))
add_para('오랜 친구이자 동료였던 자가 전쟁 중 주인공의 가문을 몰살하고, 가산과 영지를 모두 빼앗아 버린 것이다.')
add_img(img_game('cin_throne.png'), Inches(5))
add_para('주인공은 단 하나의 목적 — 복수 — 를 위해 살아간다. 마침내 배신자를 처단하고 복수를 완수하지만, 그 무게를 이기지 못하고 죽음을 맞이한다.')
add_img(img_game('cin_torture.png'), Inches(5))

doc.add_page_break()

# ACT 2 — 여신 + 지옥 + 아이들
add_para('ACT 2 — 지옥: 여신과의 계약 (Chapter 2~6)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_img(img_game('cin_nemesia_hd.png'), Inches(5))
add_para('지옥에 떨어진 주인공 앞에 복수의 여신이 나타난다. "복수가 꽤 짜릿했지? 네 놈에게 기회를 주겠다."')
add_para('여신의 실제 목적은 지옥을 손에 넣는 것이다. 주인공을 이용해 지옥왕을 제거하고, 자신은 아무런 리스크 없이 지옥 지배권을 획득하려는 계획.')
add_img(img_game('cin_remember.png'), Inches(5))
add_para('주인공은 여신의 도움으로 아이들을 구출한다. 마침내 지옥왕을 처치하고, 이 사실이 신들의 세계 전체에 알려진다.')

add_separator()

# ACT 3 — 천상
add_para('ACT 3 — 천상: 신들과의 전쟁 (Chapter 7+)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_img(img_game('cin_demonfight.png'), Inches(5))
add_para('지옥왕의 죽음으로 권력 공백이 생기자 천상의 신들이 움직인다. 주인공은 배신당했다는 사실을 알지만 여신을 용서한다. 딜의 조건: 여신이 아이들을 보호해준다. 주인공은 천상을 향해 혼자 나아간다.')
add_para('이용당하고 배신당했지만 — 여신은 주인공을 진심으로 좋아한다. 이 복잡한 감정선이 스토리의 핵심 긴장감이다.')

add_separator()
add_para('스토리 <-> 게임 시스템 연결', bold=True, size=11, color=(0x4A, 0x00, 0x00))
add_table_simple(
    ['서사 요소', '게임 시스템'],
    [
        ['복수여신의 지원', '아이템/장비/스킬 시스템의 서사적 근거'],
        ['유령 쇠뇌', '죽은 자의 원한이 깃든 고유 무기'],
        ['여신이 보낸 존재들', '펫 시스템 (까마귀, 고양이)'],
        ['여신이 계속 살려서 쓰는 구조', '죽어도 부활하는 로그라이트 루프'],
        ['인간의 악의가 지옥을 지탱', '악의 시스템 (자원/에너지)'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════
# 03. 아트 디렉션
# ════════════════════════════════════════
add_heading_styled('03. 아트 디렉션', 1)
add_para('서정적 발광 환경 x 그로테스크 적 디자인 = 잔혹동화', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
# 기존에 없는 월드 환경 컨셉 — Grok 생성분 유지
add_img(img_new('05_world.png'), Inches(5.5))
doc.add_paragraph()
add_para('01. 카라바조 명암의 지옥', bold=True, size=11)
add_para('배경은 완전한 어둠. 빛은 오직 발광 파티클과 캐릭터에서만 나온다. globalLight 0.04~0.12 + 플레이어 중심 PointLight.')
add_para('02. 인간 -> 괴물 변이 디자인', bold=True, size=11)
add_para('눈알이 무기, 입이 무기, 살덩어리가 변형한다. 베르세르크 사도급. 전형적 판타지 몬스터 (뿔 악마/드래곤) 일절 없음.')
add_para('03. 전쟁 미학의 주인공', bold=True, size=11)
add_para('작지만 실루엣 강렬. 투구/망토/대검 풀장비. 전쟁의 상흔이 디자인 언어.')
# 지옥문 이미지
add_img(img_game('hellgate.png'), Inches(4))
add_para('레퍼런스: Ori / Hollow Knight / Berserk / FromSoftware / Hades 2 / Touhou / Zelda BotW', size=9, color=(0x99, 0x99, 0x99))
doc.add_page_break()

# ════════════════════════════════════════
# 04. 핵심 캐릭터
# ════════════════════════════════════════
add_heading_styled('04. 핵심 캐릭터', 1)

# 주인공 — warrior.png (기존)
add_img(img_game('warrior.png'), Inches(3))
add_para('주인공 — 복수자 (PLR)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_para('전쟁에서 단련된 최강의 전사. 극도의 악의와 전투력을 보유해 복수의 여신에게 선택받는다. 아이들을 지키기 위해 끝없는 전투를 이어간다.')

add_separator()

# 여신 — cin_nemesia_hd.png (기존)
add_img(img_game('cin_nemesia_hd.png'), Inches(4))
add_para('복수의 여신 — 네메시아 (NEM)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_para('인간의 복수를 에너지로 삼는 존재. 냉철한 정치가이자 야심가. 주인공을 철저히 이용했지만 진심으로 그에게 끌린다. 완전한 악도 완전한 아군도 아닌, 가장 매력적인 캐릭터.')

add_separator()

# 배신자 — cin_throne.png (기존)
add_img(img_game('cin_throne.png'), Inches(4))
add_para('배신한 친구 (1막 보스)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_para('전쟁 동료이자 가장 가까운 친구였던 인물. 권력과 재산을 위해 주인공의 가문을 몰살했다. 1막의 최종 보스.')

add_para('내레이터 (NAR)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_para('길을 안내하는 존재. 주인공의 여정에 동행하며 지옥의 진실을 조금씩 드러낸다. CIN_LINES 21 라인 녹음 완료.')
doc.add_page_break()

# ════════════════════════════════════════
# 05. 월드 디자인
# ════════════════════════════════════════
add_heading_styled('05. 월드 디자인 — 7개의 지옥', 1)
# 지옥 입장 이미지
add_img(img_game('cin_fallhell_custom.png'), Inches(5))
doc.add_paragraph()
add_table_simple(
    ['장', '지옥', '고유 기믹'],
    [
        ['1', '썩은 숲', '독 웅덩이'],
        ['2', '벌레굴', '점액 감속'],
        ['3', '얼음굴', '빙판 관성'],
        ['4', '화염지대', '용암 불기둥'],
        ['5', '지옥 군단', '해골 스폰'],
        ['6', '사도의 마굴', '살점 벽'],
        ['7', '지옥성', '전 원소 트랩'],
    ]
)
doc.add_paragraph()
add_para('맵 구성', bold=True, size=11)
add_table_simple(
    ['요소', '내용'],
    [
        ['맵 타입 3종', 'Vert (세로 스네이크) / Horiz (가로 스네이크) / Grid (보스룸 원형/팔각형)'],
        ['오브젝트 5종', 'Bonfire (세이브/회복) / Lore (비문) / Breakable (파괴물) / Trap (지옥별 함정) / Altar (버프 제단)'],
        ['맵 생성', '결정론적 맵 생성 (StageSeeder)'],
        ['엔드게임', '심연의 나락 — 35 에리어 무한 리믹스, 매층 난이도 +5% 복리, 10층마다 보스, Supabase 리더보드'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════
# 06. BULLET LANGUAGE — Grok 생성분 유지 (기존에 없음)
# ════════════════════════════════════════
add_heading_styled('06. 액션 정체성 / PILLAR 1 — Bullet Language', 1)
add_para('색을 잘못 읽으면, 방어가 곧 자폭이 된다.', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
# 실제 패링 스크린샷
add_img(img_ss('스크린샷 2026-05-10 163420.png'), Inches(5.5))
doc.add_paragraph()
add_table_simple(
    ['탄막 종류', '정답 키', '정답 효과', '오답 키', '오답 효과'],
    [
        ['빨간콩 (악의의 결정체)', '우클릭 (sBash)', '블루콩 변환 / 적 추적 홈잉 / 1.5배 AoE', 'Q (sBlock)', '눈앞 폭발 / 풀뎀 / 넉백 / 슬로우모'],
        ['무지개콩 (이질적인 힘)', 'Q (sBlock)', '악의 흡수 / parryBank 충전 / 자원 회복', '우클릭 (sBash)', '눈앞 폭발 / 풀뎀 / 넉백 / 슬로우모'],
        ['일반탄 (기본기)', '우클릭 (sBash)', '반사 역투사 / 쳐낸 만큼 갚는다', '—', '오답 페널티 없음'],
    ]
)
add_para('0.3초 안에 색을 읽고, 손이 정확한 키를 눌러야 한다. 틀리면 패링이 곧 자살이 된다.', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
doc.add_page_break()

# ════════════════════════════════════════
# 07. CHAIN DRIVE — Grok 생성분 유지 (기존에 없음)
# ════════════════════════════════════════
add_heading_styled('07. 액션 정체성 / PILLAR 2 — Chain Drive', 1)
add_para('사슬 기동 — 끊기지 않는 시원한 이동과 콤보 루프', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
add_img(img_new('07_chain_drive.png'), Inches(5.5))
doc.add_paragraph()
add_para('3단계 차지 시스템', bold=True, size=11)
add_table_simple(
    ['티어', '입력', '효과'],
    [
        ['TIER 1', '탭', '기본 사거리 / 작살 발사'],
        ['TIER 2', '홀드 0.3s', '중거리 / 관통 추가'],
        ['TIER 3', '홀드 0.6s', '최장 사거리 / 무적 12프레임'],
    ]
)
doc.add_paragraph()
add_para('핵심 메커니즘', bold=True, size=11)
add_para('- 박히고 끌려간다 — 작살이 벽에 꽂히면 플레이어가 그곳으로 끌려간다. 끌려가는 동안 무적. 경로상 모든 적을 관통 히트 + 포이즈 2배.')
add_para('- 2분기 빌드 — STR 작살 (관통 데미지, 넉백 2배) / INT 사슬기동:화염 (화염 트레일, 화상 디버프)')
add_para('- 캔슬 콤보 — 회전참 복구 중 (sRecover) Shift로 즉시 캔슬 -> 작살. 4성 합체로 이어지는 무한 루프.')
doc.add_page_break()

# ════════════════════════════════════════
# 08. 스킬 & 합체 시스템
# ════════════════════════════════════════
add_heading_styled('08. 스킬 & 합체 시스템', 1)
add_para('22 액티브 + 11 패시브 + 21 합체 + 3분기 = 100+ 빌드 경로', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()
# 실제 스킬 UI 스크린샷
add_img(img_ss('스크린샷 2026-05-10 163456.png'), Inches(5))
doc.add_paragraph()
add_para('합체 쇼케이스', bold=True, size=11)
add_table_simple(
    ['조합', '합체명', '효과'],
    [
        ['회전참 + 기폭', '회전기폭', '회전하며 주변 폭발'],
        ['돌진강화 + 사슬기동', '차원돌파', '공간 관통 돌진 + 화염 트레일'],
        ['산탄 + 전방위빔 + 추적탄', '육선포', '6방향 동시 빔 (3단 합체)'],
        ['회전기폭 + 대왕치기', '슬램 스톰', '회전 + 폭발 + 내려치기 (3단)'],
        ['전 스킬 합체', '스톰 빔', '6 스킬 동시 활성 (6단 궁극)'],
    ]
)
doc.add_paragraph()
add_para('원소 x 공격 종류 시너지', bold=True, size=11)
add_table_simple(
    ['원소', '약점 무기', '효과'],
    [
        ['화상', '베기 (대검/한손검/단검)', '추가 데미지'],
        ['빙결', '둔기 (해머/철퇴/망치)', '추가 데미지'],
        ['감전', '투사체 (석궁/유탄)', '추가 데미지'],
        ['대지', '—', '적 명중률 감소'],
        ['암흑', '—', '적 회복력 감소'],
        ['신성', '—', '부활력 억제 + 언데드 추뎀'],
    ]
)
add_para('원소 반응 6종: 증기폭발 / 과냉각 / 저주불꽃 / 플라즈마 / 공허방전 / 동결심연', size=9, color=(0x99, 0x99, 0x99))
doc.add_page_break()

# ════════════════════════════════════════
# 09. 아이템 & 어픽스
# ════════════════════════════════════════
add_heading_styled('09. 아이템 & 어픽스', 1)
add_para('Path of Exile 2 식 어픽스 시스템 — 12 슬롯 x 5 등급 x 55 어픽스 풀', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()
add_para('등급 체계', bold=True, size=11)
add_table_simple(
    ['등급', '어픽스 수', '배율', '비고'],
    [
        ['일반', '0', 'x1.0', '—'],
        ['고급', '1', 'x1.2', '—'],
        ['희귀', '2', 'x1.5', '—'],
        ['영웅', '3', 'x1.8', '—'],
        ['전설', '4 + 특수', 'x2.2', 'GLOW 이펙트'],
    ]
)
doc.add_paragraph()
add_para('어픽스 풀 (55+종)', bold=True, size=11)
add_table_simple(
    ['분류', '종류', '예시'],
    [
        ['PREFIX — 공격 DoT (5)', '맹독/불꽃/냉기/번개/암흑의', '지속 피해 부여'],
        ['PREFIX — 공격 보조 (8)', '피흡/관통/광폭/처형/파쇄/연쇄/학살/폭발', '공격 효과 강화'],
        ['PREFIX — 전투 스타일 (7)', '쾌속/패리/콤보/돌진/궁극/집중/원소집중', '플레이 스타일 변경'],
        ['SUFFIX — 원소 저항 (6)', '화염/빙결/암전/암흑/독/전원소 수호', '생존 강화'],
        ['SUFFIX — 조건부 트리거 (7)', '위기/피격반격/패링폭발/스태거폭발/정화/불굴/부활', '상황별 발동'],
        ['LEGENDARY (19)', '슬롯별 고유', '전설 등급 전용 특수 효과'],
    ]
)
doc.add_paragraph()
add_para('세트 아이템 — 7 세트 (장별 고유)', bold=True, size=11)
add_para('뇌신의 권능 / 부패의 살갗 / 업화의 갑주 / 심연의 눈 / 서리의 맹세 / 백골의 왕관 / 혼돈의 파편')
add_para('2셋 / 4셋 / 6셋 보너스 구조 + 룬 소켓 1~3개 + 전설 특수 효과', size=9, color=(0x99, 0x99, 0x99))
doc.add_paragraph()
# 실제 인벤토리 UI 스크린샷
add_para('인벤토리 UI', bold=True, size=11, color=(0x8B, 0x00, 0x00))
add_img(img_ss('스크린샷 2026-05-10 163514.png'), Inches(5.5))
doc.add_page_break()

# ════════════════════════════════════════
# 10. 몬스터 디자인 — Grok 생성 몬스터 컨셉시트 유지 + bystanders
# ════════════════════════════════════════
add_heading_styled('10. 몬스터 디자인', 1)
add_para('100종 — 베르세르크 사도급 변이 디자인 / 전형적 판타지 금지', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
add_img(img_new('04_monsters.png'), Inches(5.5))
doc.add_paragraph()
add_img(img_game('cin_bystanders.png'), Inches(5))
doc.add_paragraph()
add_para('디자인 원칙', bold=True, size=11)
add_para('01. 눈알이 무기다 — 몸 곳곳의 눈에서 탄막을 쏜다')
add_para('02. 입이 무기다 — 거대한 입에서 레이저와 포식 공격')
add_para('03. 살덩어리가 변형한다 — 팔이 칼날, 등이 촉수, 배가 입')
add_para('04. 인간의 흔적이 남아있다 — 얼굴, 손, 비명 자세')
add_para('05. 크기 차이가 극단적 — 쥐만한 기생충부터 화면 가득 거인')
doc.add_paragraph()
add_table_simple(
    ['분류', '수량', '설명'],
    [
        ['공통', '20종', '전 지옥 출현'],
        ['지옥전용', '70종', '지옥별 10종 테마 몹'],
        ['희귀', '10종', '방랑기사, 상인악마, 죽음 그 자체'],
    ]
)
add_para('금지: 뿔 악마, 날개 드래곤, 갑옷 해골 병사, 색만 바꾼 슬라임, 그냥 큰 늑대/거미', size=9, color=(0xCC, 0x00, 0x00))
doc.add_page_break()

# ════════════════════════════════════════
# 11. 보스 시스템 — cin_demonbattle (기존)
# ════════════════════════════════════════
add_heading_styled('11. 보스 시스템', 1)
add_para('19 보스 (12 미니 + 6 장보스 + 1 최종) / 48 패턴 / 58 콤보 / 6 아키타입', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
add_img(img_game('cin_demonbattle.png'), Inches(5.5))
doc.add_paragraph()
add_para('6 아키타입', bold=True, size=11)
add_table_simple(
    ['코드', '유형', '레퍼런스'],
    [
        ['A1', '근접 광전사', '베르세르크 구츠 / 말레니아'],
        ['A2', '중장갑 폭군', '오르스타인 / 라오샨룽'],
        ['A3', '원거리 마법사', '아스텔 / 그림'],
        ['A4', '다리걸림 거수', '섀도우 오브 콜로서스 1'],
        ['A5', '탄막 커튼', '동방 / 라다곤 유성'],
        ['A6', '환경 조작', '벨카 / 세키로 사자원숭이'],
    ]
)
doc.add_paragraph()
add_para('부활 시스템', bold=True, size=11)
add_para('- 보장 부활 = 지옥 단수 (1단=1회, 7단=7회) + 확률 부활')
add_para('- 부활은 단순 HP 리셋이 아닌 "새 페이즈". 외형 변화 -> 지형 변화 -> 최후 시네마틱.')
add_para('- 플레이어 카운터: antiRevive 어픽스 + holyPrison 신성력 존으로 억제 가능.')
add_para('- 패링 색 코드: 빨간색 = 패링 가능 (PARRY 링) / 무지개색 = 패링 불가 (DANGER 링)')
doc.add_page_break()

# ════════════════════════════════════════
# 12. 기술 아키텍처
# ════════════════════════════════════════
add_heading_styled('12. 기술 아키텍처', 1)
add_para('단일 game.html 30,000+ 줄 / 솔로 인디로서 독보적 기술 규모', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()
add_table_simple(
    ['영역', '기술 스택'],
    [
        ['RENDERING', 'WebGL2 (ANGLE -> D3D11) + Canvas 2D 혼합\nProxyX 배칭 시스템\n파티클 풀링 1000개 + VFX 소프트캡 150\n오프스크린 프리렌더 (10,648 -> 1 drawCall)'],
        ['SIMULATION', '결정론적 맵 생성 (StageSeeder)\n청크 파티셔닝 + 공간 해시 (shQuery)\n플로우필드 패스파인딩 (Web Worker)\n오브젝트 풀링 — 게임루프 내 new/splice 금지'],
        ['SYSTEMS', '사운드 우선순위 시스템 (동시 40+ 재생)\n보스 콤보 idx 48 패턴 / 58 콤보\n6 원소 x 6 반응 조합 매트릭스\nSupabase — 리더보드, 아이템 검증'],
        ['SHIPPING', 'Electron 번들 (Steam 타겟)\nNode.exe 번들 ZIP (오프라인 배포)\nHTML5 데모 (Steam 위시리스트 연동)\ngit 자동 커밋 (Windows 작업 스케줄러)'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════
# 13. AI-First 파이프라인
# ════════════════════════════════════════
add_heading_styled('13. AI-First 제작 파이프라인', 1)
add_para('1인 개발자가 AAA 규모 콘텐츠를 생산하는 실제 파이프라인', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()
add_table_simple(
    ['영역', '도구', '용도'],
    [
        ['CODE', 'Claude Code (Opus) + VS Code', '30K+ 줄 단일 파일 유지보수'],
        ['SPRITE', 'PixelLab MCP + God Mode AI', '캐릭터/몬스터 스프라이트 + 애니메이션'],
        ['ILLUSTRATION', 'ComfyUI + Berserk LoRA + Midjourney', '일러스트, 맵 배경, 보스 아트'],
        ['3D / MOTION', 'Blender MCP + Mixamo', '장보스 3D 모션 -> PNG 시퀀스 변환'],
        ['SOUND', 'ElevenLabs + Freesound.org', 'SFX 생성, 내레이션 21 라인, 루프 BGM'],
        ['BACKEND', 'Supabase', '리더보드, 아이템 검증, 비동기 멀티플레이'],
        ['ICON', 'Python + PIL + rembg (u2net)', '157 아이콘 슬라이싱 자동화, 13 카테고리'],
    ]
)
add_para('결과: 솔로 개발자 1명이 7 챕터 35 스테이지 100 몬스터 55 어픽스 19 보스 규모를 실제로 구현 가능', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x4A, 0x00, 0x00))
doc.add_page_break()

# ════════════════════════════════════════
# 14. 개발 현황
# ════════════════════════════════════════
add_heading_styled('14. 개발 현황', 1)
add_para('뼈대 완성 — 아트/사운드 조립 단계. 완성 리스크 낮은 프로젝트.', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()
add_para('DONE — 완성', bold=True, size=11, color=(0x00, 0x80, 0x00))
for item in [
    '7 챕터 35 스테이지 게임 프레임',
    '전투 / 패링 / 콤보 시스템 완전 구현',
    '48 보스 패턴 / 58 콤보 코드',
    '22 액티브 + 11 패시브 + 21 합체 스킬',
    '55+ 어픽스 풀 + 전설 특수 효과 19',
    '100 몬스터 타입 + 28 엘리트 모디파이어',
    '인트로 시네마틱 + 내레이션 21 라인',
    '240fps 안정 구동 확인',
]:
    add_para(f'  V  {item}', size=10)

doc.add_paragraph()
add_para('IN PROGRESS — 진행 중', bold=True, size=11, color=(0xCC, 0x88, 0x00))
for item in [
    '보스전 비주얼 / 연출 디자인 (P0 3마리)',
    '아트 에셋 조립 (오리 분위기 맵)',
    '사운드 디자인 (젤다 레퍼런스 레이어드 SFX)',
    'MAP 최적화 Phase 6 (프리렌더 + AABB 컬링)',
    '보스 모션 파이프라인 (Blender MCP + Mixamo)',
    '대장간 UI 탭 시스템',
    '로컬라이제이션 KR -> EN -> CN -> JP',
    'BIC 2026 제출 빌드 패키징',
]:
    add_para(f'  O  {item}', size=10)

add_para('보스 구현 전략: 7 장보스 독창 (2주/마리) + 14 변형 (3일/마리) + 14 엘리트몹 승격 (1일/마리) = 5개월', size=9, color=(0x99, 0x99, 0x99))
doc.add_page_break()

# ════════════════════════════════════════
# 15. 출시 로드맵
# ════════════════════════════════════════
add_heading_styled('15. 출시 로드맵', 1)
add_para('BIC 2026 -> Neowiz 퍼블리싱 피치 -> Steam Early Access -> 풀릴리즈 -> DLC', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()
add_table_simple(
    ['Phase', '시기', '내용'],
    [
        ['PHASE 1', '2026 상반기', 'P0 보스 3마리 완성\nBIC 부산인디커넥트 제출\nIndieGo 2026 스마일게이트 제출\nSteam 스토어 페이지 오픈'],
        ['PHASE 2', '2026 하반기', 'YouTube/Twitch 개발 방송\n스트리머 타임어택 이벤트\nSteam 위시리스트 오픈\nNeowiz 퍼블리싱 협상'],
        ['PHASE 3', '2027', '퍼블리셔 파트너십 체결\nEarly Access 출시 ($18-20)\n투자 유치 및 인력 채용\nASH RAIDER DLC 프로덕션 착수'],
    ]
)
add_para('벤치마크: Shape of Dreams (Neowiz 퍼블리싱, 4인팀, 100만 카피). 솔로 + AI 파이프라인으로 동급 스케일 도달.', size=9, color=(0x99, 0x99, 0x99))
doc.add_page_break()

# ════════════════════════════════════════
# 16. DLC & 포스트런칭
# ════════════════════════════════════════
add_heading_styled('16. DLC & 포스트런칭', 1)
add_para('DLC 01 — ASH RAIDER (재의 딸)', bold=True, size=12, color=(0x8B, 0x00, 0x00))
add_para('여성 주인공 / 대도끼 + 의수 팔 / 본편 세계관 연결. 본편 클리어 이후 해금되는 독립 캐릭터. 완전히 다른 무기 철학과 스킬 트리. 새로운 지옥 지역 + 신규 보스 5마리 + 신규 엔딩 분기.')
doc.add_paragraph()
add_para('무료 업데이트 로드맵', bold=True, size=11)
add_table_simple(
    ['버전', '제목', '내용'],
    [
        ['v1.1', '신규 엔드게임', '심연의 나락 추가 층수 + 새 모디파이어'],
        ['v1.2', 'PvP 고스트', '비동기 리플레이 + 타임어택 경쟁'],
        ['v1.3', '협동 P2P', 'PeerJS 기반 2인 협동 (BIC 이후 로드맵)'],
    ]
)
doc.add_paragraph()
add_table_simple(
    ['DLC', '내용', '가격'],
    [
        ['DLC 1 — ASH RAIDER', '신 캐릭터 + 신 지역 + 신규 보스 5', '$10-12'],
        ['DLC 2 — 심연의 여왕', 'TBA — 본편 엔딩 분기 확장', 'TBA'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════
# 마지막 페이지
# ════════════════════════════════════════
doc.add_paragraph()
add_img(img_game('etype84_lightning.png'), Inches(4))
doc.add_paragraph()
add_para('아름다운 지옥 위에서 기괴한 괴물을 쓸어죽이는 게임', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x4A, 0x00, 0x00))
doc.add_paragraph()
add_para('BULLET LANGUAGE  /  CHAIN DRIVE  /  HELLSWARM', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x8B, 0x00, 0x00))
doc.add_paragraph()
add_para('복수를 끝낸 남자가 지옥에 떨어진다.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('그리고 더 깊은 이야기가 시작된다.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_img(img_ss('스크린샷 2026-05-10 163644.png'), Inches(4))
doc.add_paragraph()
add_para('GAME DESIGN DOCUMENT v4.1  |  2026', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))

doc.save(OUT_PATH)
print(f'Saved: {OUT_PATH}')
